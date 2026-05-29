from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

h2 = doc.add_heading('4.2 README Dosyasi', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 102, 153)

p_inst = doc.add_paragraph()
run_inst = p_inst.add_run('[README dosyanizin icerigini ozetleyiniz. README; proje aciklamasi, kurulum adimlari, calistirma talimatlari, kullanilan teknolojiler, ekran goruntuleri ve takim uyeleri bolumlerini icermelidir.]')
run_inst.italic = True
run_inst.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('Projemizin ana dizininde yer alan README.md dosyasi, uygulamanin gelistiriciler ve kullanicilar tarafindan anlasilmasini saglayan temel basvuru dokumanimizdir. Dosya icerigi asagidaki bolumlerden olusmaktadir:')

def add_bold_bullet(title, text):
    p = doc.add_paragraph(style='List Bullet')
    run_title = p.add_run(title)
    run_title.font.bold = True
    p.add_run(' ' + text)

add_bold_bullet('1. Proje Aciklamasi:', 'eMuhasebe Pro; kucuk ve orta olcekli isletmeler (veya okul projesi kapsami) icin gelistirilmis, yapay zeka destekli ve gercek zamanli (realtime) senkronizasyon ozelliklerine sahip gelismis bir on muhasebe yonetim sistemidir.')

add_bold_bullet('2. Kullanilan Teknolojiler:', 'Backend: Python 3.12+, Flask, SQLAlchemy, Flask-JWT-Extended, Pydantic. Frontend: Vanilla JS, HTML5 (Jinja2), CSS3. Veritabani: Firebase Realtime Database, SQLite. Yapay Zeka: Google Gemini API.')

add_bold_bullet('3. Kurulum Adimlari:', 'Projenin bilgisayara indirilmesinden sonra bir Python sanal ortami (.venv) olusturulmasi ve requirements.txt dosyasindaki gerekli kutuphanelerin (pip install -r requirements.txt) yuklenmesi sureci anlatilmistir.')

add_bold_bullet('4. Calistirma Talimatlari:', 'Sistemin calisabilmesi icin gereken ortam degiskenlerinin (.env dosyasi) nasil ayarlanacagi belirtilmistir. Veritabani tablolarinin olusturulmasi (flask db upgrade) ve yerel sunucunun baslatilmasi (python run.py) komutlari detaylandirilmistir.')

add_bold_bullet('5. Takim Uyeleri:', '[Bu kisma sizin ve eger varsa proje arkadaslarinizin isimlerini ekleyiniz.]')

add_bold_bullet('6. Ekran Goruntuleri:', '[Raporda bu yazinin hemen altina eMuhasebe Pro nun 1-2 adet ekran goruntusunu ekleyiniz.]')

save_path = r'C:\Users\rua\OneDrive\Desktop\README_Raporu.docx'
doc.save(save_path)
