from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Varsayılan stil
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ========== BAŞLIK ==========
h1 = doc.add_heading('11. SONUÇ VE DEĞERLENDİRME', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in h1.runs:
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

# ========== 11.1 ==========
h2 = doc.add_heading('11.1 Elde Edilen Kazanımlar', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph('eMuhasebe Pro projesinin geliştirilme sürecinde hem teknik hem de takım çalışması açısından önemli kazanımlar elde edilmiştir:')

h3 = doc.add_heading('Teknik Kazanımlar', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

teknik = [
    ('ORM ve Repository Pattern:', 'SQLAlchemy ORM kullanılarak veritabanı işlemleri nesne yönelimli bir yapıya taşınmıştır. Repository Pattern ile veri erişim katmanı, iş mantığı katmanından tamamen ayrılmış; böylece kod tekrarı önlenmiş ve test edilebilirlik artırılmıştır. BaseRepository sınıfından türeyen MusteriRepository, UrunRepository gibi alt sınıflar sayesinde CRUD işlemleri merkezi bir yapıda yönetilmiştir.'),
    ('Katmanlı Mimari (Layered Architecture):', 'Proje; Model → Repository → Service → API → Template şeklinde beş katmanlı bir mimari üzerine inşa edilmiştir. Her katmanın sorumluluğu net bir şekilde ayrılmış, bu sayede bir katmandaki değişikliğin diğer katmanları etkilemesi minimize edilmiştir.'),
    ('RESTful API Tasarımı:', 'Flask blueprint yapısıyla modüler REST API endpoint\'leri tasarlanmıştır. Pydantic ile güçlü giriş validasyonu, JWT ile kimlik doğrulama, Flask-Limiter ile rate limiting gibi endüstri standartlarında güvenlik katmanları uygulanmıştır.'),
    ('Unit Test ve Test Kültürü:', 'pytest framework\'ü ile 11 farklı test modülü yazılmıştır. Repository, service, API, validasyon ve güvenlik katmanlarının her biri için ayrı test dosyaları oluşturulmuş; conftest.py ile ortak test fixture\'ları tanımlanmıştır. Bu süreçte test yazma disiplini ve test-driven düşünme alışkanlığı kazanılmıştır.'),
    ('CI/CD Pipeline:', 'GitHub Actions ile otomatik sürekli entegrasyon pipeline\'ı kurulmuştur. Her commit ve pull request\'te otomatik olarak flake8 ile kod kalite kontrolü, pytest ile birim testleri çalıştırılmakta ve başarılı build sonrası otomatik deploy tetiklenmektedir.'),
    ('Firebase Realtime Database Entegrasyonu:', 'Uygulamanın veritabanı katmanı, hem yerel tarayıcı depolama (LocalStorage) hem de bulut tabanlı Firebase Realtime Database ile çalışacak şekilde tasarlanmıştır. Bu dual-mode mimari sayesinde kullanıcılar internet bağlantısı olmadan da uygulamayı kullanabilmekte, bağlantı geldiğinde veriler otomatik senkronize edilmektedir.'),
    ('Yapay Zeka (AI) Entegrasyonu:', 'Google Gemini AI API\'si backend\'e entegre edilerek, kullanıcının finansal verilerini analiz eden ve Türkçe öneriler sunan bir yapay zeka asistanı geliştirilmiştir. Prompt engineering teknikleriyle AI\'dan yapılandırılmış ve anlamlı çıktılar alınması sağlanmıştır.'),
    ('Frontend Tasarım ve UX:', 'Modern CSS teknikleri (CSS Custom Properties, Glassmorphism, CSS Grid/Flexbox), responsive tasarım, koyu tema (dark mode), mikro animasyonlar ve erişilebilirlik (ARIA etiketleri) gibi güncel web standartları uygulanarak profesyonel bir kullanıcı deneyimi oluşturulmuştur.'),
]

for bold_part, text_part in teknik:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_text = p.add_run(text_part)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

h3 = doc.add_heading('Takım Çalışması Kazanımları', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

takim = [
    ('Git/GitHub İş Akışı:', 'Feature branch modeli (feature/MHSB-*) uygulanmış, her yeni özellik için ayrı dal oluşturulmuş ve Pull Request süreciyle main branch\'e merge edilmiştir. Bu sayede paralel geliştirme yapabilme ve versiyon kontrolü disiplini edinilmiştir.'),
    ('Code Review Kültürü:', 'Her Pull Request\'te en az bir takım üyesinin inceleme (review) yapması zorunlu tutulmuştur. Bu süreçte kod okunabilirliği, hata tespiti ve bilgi paylaşımı konularında önemli deneyim kazanılmıştır.'),
    ('Kanban ve Proje Yönetimi:', 'MHSB numaralı görevlerle (issue) takip edilen bir Kanban akışı kullanılmıştır. Görevler "To Do → In Progress → Review → Done" aşamalarından geçirilerek süreç yönetimi disiplini uygulanmıştır.'),
]

for bold_part, text_part in takim:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_text = p.add_run(text_part)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

# ========== 11.2 ==========
h2 = doc.add_heading('11.2 Karşılaşılan Zorluklar ve Çözümleri', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

zorluklar = [
    ('1', 'Merge Çakışmaları: Birden fazla feature branch\'in aynı dosyaları düzenlemesi sonucu merge conflict\'ler yaşanmıştır. Özellikle app/__init__.py ve models.py dosyalarında çakışmalar sık karşılaşılmıştır.',
     'Git\'in merge araçları ve manuel çakışma çözümü kullanılmıştır. Çakışma yaşanan her dosya satır satır incelenmiş, doğru versiyon korunarak temiz bir merge gerçekleştirilmiştir. İlerleyen süreçte dosya düzenleme sorumluluklarının branch bazında ayrılması ile çakışma sıklığı azaltılmıştır.'),
    ('2', 'Firebase ve LocalStorage Arasında Veri Senkronizasyonu: Uygulamanın hem çevrimdışı (LocalStorage) hem de çevrimiçi (Firebase) modda çalışması, veri tutarlılığı açısından karmaşıklık yaratmıştır.',
     'Dual-mode bir veritabanı soyutlama katmanı (db.js) geliştirilmiştir. getMode() fonksiyonu ile aktif mod tespit edilmekte, tüm CRUD işlemleri bu katman üzerinden yürütülmektedir. Mod geçişlerinde mevcut verilerin yeni ortama taşınması için senkronizasyon mekanizması eklenmiştir.'),
    ('3', 'API Anahtarlarının Güvenliği: Firebase ve AI API anahtarlarının kaynak koduna gömülmesi güvenlik riski oluşturuyordu.',
     '.env dosyası ile ortam değişkenleri kullanılmış, .gitignore ile bu dosyanın GitHub\'a yüklenmesi engellenmiştir. Backend tarafında /api/config/firebase endpoint\'i ile anahtarlar güvenli şekilde frontend\'e aktarılmıştır. .env.example dosyası ile diğer geliştiricilere hangi anahtarların gerektiği bildirilmiştir.'),
    ('4', 'Fatura Hesaplama Mantığının Karmaşıklığı: Farklı KDV oranları (%0, %1, %8, %10, %18, %20), kalem bazlı hesaplama, iskonto ve genel toplam gibi çoklu hesaplamaların doğru çalışması zorlu olmuştur.',
     'Her hesaplama adımı ayrı fonksiyonlara bölünmüş ve test_hesaplama.py dosyasında kapsamlı birim testler yazılmıştır. Farklı KDV senaryoları, sıfır tutarlı faturalar ve negatif değer kontrolleri test edilmiştir.'),
    ('5', 'Yapay Zeka Yanıtlarının Formatlanması: AI\'dan gelen yanıtlar düz metin olarak geldiğinde, özellikle numaralı listeler ve başlıklar tarayıcıda okunamaz hale geliyordu.',
     'Özel bir formatMarkdown() JavaScript fonksiyonu yazılmıştır. Bu fonksiyon; bold metinleri, numaralı listeleri, madde işaretlerini ve paragraf ayrımlarını HTML\'e dönüştürerek AI yanıtlarının profesyonel ve okunabilir biçimde görüntülenmesini sağlamıştır.'),
    ('6', 'Türkçe Karakter Problemleri: Loglama ve konsol çıktılarında Türkçe karakterlerin (ğ, ü, ş, ı, ö, ç) bozuk görünmesi sorunu yaşanmıştır.',
     'PYTHONIOENCODING=utf-8 ortam değişkeni ayarlanmış, tüm dosyalar UTF-8 encoding ile kaydedilmiş ve HTML meta etiketlerinde charset="UTF-8" tanımlanmıştır.'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
hdr = table.rows[0].cells
for i, text in enumerate(['#', 'Zorluk', 'Çözüm']):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    # Arka plan rengi
    shading = hdr[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'D9E2F3'
    })
    shading.append(shading_elm)

# Sütun genişlikleri
table.columns[0].width = Cm(1)
table.columns[1].width = Cm(7)
table.columns[2].width = Cm(8.5)

for num, zorluk, cozum in zorluklar:
    row = table.add_row().cells
    for i, text in enumerate([num, zorluk, cozum]):
        row[i].text = ''
        p = row[i].paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

# ========== 11.3 ==========
h2 = doc.add_heading('11.3 Gelecek Çalışmalar', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph('Projenin mevcut hali, tam işlevsel bir ön muhasebe sistemi olarak kullanılabilir durumdadır. Ancak aşağıdaki iyileştirmeler ile uygulamanın kapsamı ve kalitesi daha da artırılabilir:')

h3 = doc.add_heading('Yeni Özellikler', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

yeni = [
    ('Kasa ve Banka Yönetimi:', 'Tahsilat ve ödeme işlemlerinin takip edilebileceği bir finans modülü eklenebilir. Fatura ödemeleri kasa/banka hareketleriyle eşleştirilerek nakit akışı raporları oluşturulabilir.'),
    ('Cari Hesap Ekstresi:', 'Müşteri ve tedarikçi bazında detaylı cari hesap ekstresi oluşturulabilir. Borç-alacak dengesi, vadesi geçmiş alacaklar ve ödeme geçmişi takip edilebilir.'),
    ('PDF Rapor Oluşturma:', 'Dashboard\'daki finansal özetler, faturalar ve AI analiz sonuçları PDF formatında indirilebilir hale getirilebilir.'),
    ('Stok Takibi:', 'Ürün giriş-çıkış hareketleri ile stok miktarı otomatik güncellenebilir. Minimum stok seviyesi uyarıları ve stok değerleme raporları eklenebilir.'),
]

for bold_part, text_part in yeni:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_text = p.add_run(text_part)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

h3 = doc.add_heading('Yapay Zeka İyileştirmeleri', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

ai_items = [
    ('Anomali Tespiti:', 'AI analizi genişletilerek olağandışı harcamalar, beklenmedik fatura artışları ve şüpheli işlemler otomatik tespit edilebilir.'),
    ('Tahminleme (Forecasting):', 'Geçmiş veriler analiz edilerek gelecek ay/çeyrek için gelir-gider tahminleri yapılabilir.'),
    ('Doğal Dil ile Sorgulama:', 'Kullanıcının "Bu ay ne kadar harcadık?" gibi doğal dil sorularıyla veri sorgulayabilmesi sağlanabilir.'),
]

for bold_part, text_part in ai_items:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_text = p.add_run(text_part)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

h3 = doc.add_heading('Performans ve Altyapı İyileştirmeleri', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

perf = [
    ('Progressive Web App (PWA):', 'Mevcut Service Worker altyapısı genişletilerek tam çevrimdışı çalışma desteği ve push bildirimleri eklenebilir.'),
    ('Mobil Uygulama:', 'Flutter veya React Native ile iOS ve Android platformlarına özel bir mobil sürüm geliştirilebilir.'),
    ('Çoklu Kullanıcı ve Rol Yönetimi:', 'Şu anda tek kullanıcılı olan sistem, "Yönetici", "Muhasebeci", "Görüntüleyici" gibi rol tabanlı yetkilendirme sistemiyle çoklu kullanıcı desteğine kavuşturulabilir.'),
    ('E-Fatura Entegrasyonu:', 'GİB (Gelir İdaresi Başkanlığı) e-Fatura ve e-Arşiv Fatura servisleriyle entegrasyon sağlanarak yasal uyumluluk artırılabilir.'),
    ('Veritabanı Optimizasyonu:', 'Büyüyen veri hacmine karşı indeksleme, sayfalama (pagination) ve önbellekleme (caching) stratejileri uygulanabilir.'),
]

for bold_part, text_part in perf:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_text = p.add_run(text_part)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

# Kaydet
doc.save(r'c:\Users\rua\OneDrive\Desktop\Bolum11_Sonuc_ve_Degerlendirme.docx')
print('Word dosyasi basariyla olusturuldu!')
