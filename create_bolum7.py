"""
Bölüm 7 – Uygulama Geliştirme (Word belgesi oluşturma)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Stil Ayarları ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Sayfa kenar boşlukları ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def set_cell_shading(cell, color):
    """Hücre arka plan rengini ayarla."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_data, is_header=False):
    if is_header:
        row = table.rows[0]
    else:
        row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if is_header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '334155')
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row


def add_code_block(code_text, description=""):
    """Kod bloğu ekle (gri arka plan simülasyonu)."""
    if description:
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True

    # Kod için tablo kullan (arka plan rengi için)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F1F5F9')
    cell.text = ''
    for line in code_text.strip().split('\n'):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)
    # İlk boş paragrafı sil
    if cell.paragraphs[0].text == '':
        p_element = cell.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    doc.add_paragraph()  # boşluk


# ═══════════════════════════════════════════════════
# BÖLÜM 7 – UYGULAMA GELİŞTİRME
# ═══════════════════════════════════════════════════

add_heading_styled('7. UYGULAMA GELİŞTİRME', level=1)

# ────────────────────────────
# 7.1 Modüller ve Bileşenler
# ────────────────────────────

add_heading_styled('7.1 Modüller ve Bileşenler', level=2)

doc.add_paragraph(
    'eMuhasebe Pro uygulaması, katmanlı mimari (Layered Architecture) prensibi ile '
    'tasarlanmış olup her modül kendi sorumluluk alanına sahiptir. '
    'Aşağıda uygulamayı oluşturan ana modüller, bileşenler ve aralarındaki bağımlılıklar açıklanmaktadır.'
)

# Modül tablosu
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Tablo 7.1 – Ana Modüller ve Sorumlulukları')
run.bold = True
run.font.size = Pt(10)

t1 = doc.add_table(rows=1, cols=3)
t1.style = 'Table Grid'
add_table_row(t1, ['Modül', 'Konum', 'Açıklama'], is_header=True)

modules = [
    ['Auth Module\n(Kimlik Doğrulama)',
     'app/auth/',
     'JWT tabanlı kullanıcı kaydı, giriş, çıkış ve token yenileme işlemlerini yönetir. '
     'Redis veya in-memory token blocklist desteği sunar. '
     'Flask-JWT-Extended kütüphanesi kullanılır.'],
    ['Müşteri Module\n(Customer)',
     'app/musteriler/\napp/services/musteri_service.py\napp/repositories/musteri_repository.py',
     'Müşteri/tedarikçi CRUD işlemleri. Vergi numarası ve e-posta benzersizlik kontrolü, '
     'arama ve filtreleme özelliği sunar. Repository Pattern ile veri erişimi sağlanır.'],
    ['Ürün Module\n(Product)',
     'app/urunler/\napp/services/urun_service.py\napp/repositories/urun_repository.py',
     'Ürün/hizmet tanımları, stok takibi, KDV oranı yönetimi. '
     'Numeric(12,2) hassasiyetiyle finansal doğruluk sağlanır.'],
    ['Fatura Module\n(Invoice)',
     'app/faturalar/\napp/services/fatura_service.py',
     'Alış, satış ve iade faturası yönetimi. Her fatura türü için ayrı repository ve blueprint. '
     'Fatura hesaplama motoru indirim, KDV matrahı ve genel toplamı otomatik hesaplar.'],
    ['API Module\n(REST API)',
     'app/api/__init__.py\napp/api/v1/',
     'Tüm REST API endpoint\'lerini barındırır. Rate limiting (Flask-Limiter), '
     'OpenAPI 3.0 spec (/openapi.json) ve healthcheck desteği sunar.'],
    ['AI Module\n(Yapay Zeka)',
     'app/static/js/ai-engine.js\napp/api/__init__.py (/ai/*)',
     'Üç bağımsız yapay zeka modülü: (1) Nakit akışı tahmini – lineer regresyon + '
     'hareketli ortalama, (2) Anomali tespiti – Z-score tabanlı, '
     '(3) Akıllı fatura önerisi – geçmiş faturalardan ürün/fiyat önerisi. '
     'Backend tarafında Gemini / OpenRouter API proxy desteği.'],
    ['Firebase Module\n(Veritabanı)',
     'app/static/js/db.js\napp/static/js/firebase-db.js',
     'Firebase Realtime Database entegrasyonu. Müşteri, ürün ve fatura verilerini '
     'gerçek zamanlı (realtime) olarak okur/yazar. LocalStorage fallback mekanizması ile '
     'Firebase yapılandırması yoksa çevrimdışı çalışabilir.'],
    ['Validator Module\n(Doğrulama)',
     'app/validators.py',
     'Pydantic tabanlı API input doğrulama. Her varlık (müşteri, ürün, fatura) için '
     'ayrı request/response modelleri. KDV oranı, birim, tarih ve fiyat formatı kontrolü.'],
    ['Model Module\n(Veritabanı Modelleri)',
     'app/models.py',
     'SQLAlchemy ORM modelleri. SoftDeleteMixin ile mantıksal silme, AuditLog ile '
     'tüm değişikliklerin izlenmesi. CHECK constraint\'ler ile veri bütünlüğü.'],
    ['Middleware & Logging',
     'app/middleware.py\napp/logging_config.py',
     'İstek/yanıt loglama, AuditLogContext context manager, yapılandırılmış '
     'loglama sistemi (dosya + konsol).'],
]

for mod in modules:
    add_table_row(t1, mod)

# Modül bağımlılık açıklaması
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Modüller Arası Bağımlılık İlişkileri:')
run.bold = True
run.font.size = Pt(11)

deps = [
    'API Module → Service Layer → Repository Layer → Model Layer (SQLAlchemy ORM) → Veritabanı (SQLite/PostgreSQL)',
    'Auth Module → User Model → JWT Manager → Token Blocklist (Redis / In-Memory)',
    'Firebase Module → Firebase Realtime Database → LocalStorage Fallback',
    'AI Module (Frontend) → AI Engine (JavaScript) | AI Module (Backend) → Gemini API / OpenRouter API',
    'Validator Module → Pydantic BaseModel → API Endpoints (Request validation)',
    'Middleware → Flask before_request / after_request → AuditLog Model',
]

for dep in deps:
    p = doc.add_paragraph(dep, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_paragraph()

# ────────────────────────────
# 7.2 API / Servis Katmanı
# ────────────────────────────

add_heading_styled('7.2 API / Servis Katmanı', level=2)

doc.add_paragraph(
    'eMuhasebe Pro, RESTful API mimarisi üzerine kurulmuştur. API endpoint\'leri Flask Blueprint yapısıyla '
    'modüler olarak düzenlenmiştir. Tüm endpoint\'ler rate limiting ile korunmakta ve '
    'OpenAPI 3.0 spesifikasyonu sunmaktadır. Aşağıdaki tabloda tüm API uç noktaları listelenmiştir.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Tablo 7.2 – REST API Endpoint\'leri')
run.bold = True
run.font.size = Pt(10)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
add_table_row(t2, ['Method', 'Endpoint', 'Yetki', 'Açıklama'], is_header=True)

endpoints = [
    # Sistem
    ['GET', '/api/health', 'Anonim', 'Sistem sağlık kontrolü (veritabanı bağlantı durumu).'],
    ['GET', '/api/openapi.json', 'Anonim', 'OpenAPI 3.0 spesifikasyon belgesi.'],
    ['GET', '/api/config/firebase', 'Anonim', 'Firebase yapılandırmasını ortam değişkenlerinden döner.'],
    # Auth
    ['POST', '/api/v1/auth/register', 'Anonim*', 'Yeni kullanıcı kaydı oluşturur. İlk kullanıcı dışındakiler admin yetkisi gerektirir.'],
    ['POST', '/api/v1/auth/login', 'Anonim', 'E-posta ve şifre ile giriş yapar; JWT access ve refresh token döner.'],
    ['POST', '/api/v1/auth/refresh', 'JWT (Refresh)', 'Refresh token ile yeni access token alır.'],
    ['GET', '/api/v1/auth/me', 'JWT', 'Aktif kullanıcı bilgilerini döner.'],
    ['POST', '/api/v1/auth/logout', 'JWT', 'Token\'ı blocklist\'e ekleyerek oturumu sonlandırır.'],
    # Müşteriler
    ['GET', '/api/musteriler', 'Tüm Kullanıcılar', 'Müşteri listesini döner. ?q= parametresi ile arama destekler.'],
    ['GET', '/api/musteriler/<id>', 'Tüm Kullanıcılar', 'Belirtilen ID\'ye sahip müşterinin detayını döner.'],
    ['POST', '/api/musteriler', 'Tüm Kullanıcılar', 'Yeni müşteri oluşturur (ünvan, vergi no, e-posta vb.).'],
    ['PUT', '/api/musteriler/<id>', 'Tüm Kullanıcılar', 'Müşteri bilgilerini günceller.'],
    ['DELETE', '/api/musteriler/<id>', 'Tüm Kullanıcılar', 'Müşteriyi mantıksal olarak siler (soft delete).'],
    # Ürünler
    ['GET', '/api/urunler', 'Tüm Kullanıcılar', 'Ürün listesini döner. ?q= parametresi ile arama destekler.'],
    ['GET', '/api/urunler/<id>', 'Tüm Kullanıcılar', 'Belirtilen ID\'ye sahip ürünün detayını döner.'],
    ['POST', '/api/urunler', 'Tüm Kullanıcılar', 'Yeni ürün/hizmet oluşturur (kod, ad, fiyat, KDV oranı).'],
    ['PUT', '/api/urunler/<id>', 'Tüm Kullanıcılar', 'Ürün bilgilerini günceller.'],
    ['DELETE', '/api/urunler/<id>', 'Tüm Kullanıcılar', 'Ürünü mantıksal olarak siler (soft delete).'],
    # Faturalar
    ['GET', '/api/faturalar/ozet', 'Tüm Kullanıcılar', 'Dashboard için fatura özet istatistiklerini döner.'],
    ['GET', '/api/faturalar/alis', 'Tüm Kullanıcılar', 'Tüm alış faturalarını listeler.'],
    ['POST', '/api/faturalar/alis', 'Tüm Kullanıcılar', 'Yeni alış faturası oluşturur (kalemler dahil).'],
    ['GET', '/api/faturalar/satis', 'Tüm Kullanıcılar', 'Tüm satış faturalarını listeler.'],
    ['POST', '/api/faturalar/satis', 'Tüm Kullanıcılar', 'Yeni satış faturası oluşturur (kalemler dahil).'],
    ['GET', '/api/faturalar/iade', 'Tüm Kullanıcılar', 'Tüm iade faturalarını listeler.'],
    ['POST', '/api/faturalar/iade', 'Tüm Kullanıcılar', 'Yeni iade faturası oluşturur (kalemler dahil).'],
    # AI
    ['POST', '/api/ai/ping', 'Tüm Kullanıcılar', 'AI API anahtarını doğrular (Gemini veya OpenRouter).'],
    ['POST', '/api/ai/analyze', 'Tüm Kullanıcılar', 'Yapay zeka analiz endpoint\'i; finansal veri analizi yapar.'],
]

for ep in endpoints:
    add_table_row(t2, ep)

# Sütun genişliklerini ayarla
for row in t2.rows:
    row.cells[0].width = Cm(2)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(3)
    row.cells[3].width = Cm(7)

# Örnek istek/yanıt
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Örnek API İstek ve Yanıtları:')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Örnek 1 – Müşteri Oluşturma (POST /api/musteriler)')
run.bold = True
run.font.size = Pt(10)

add_code_block('''// İstek (Request Body)
{
    "unvan": "ABC Ltd Şti",
    "vergi_no": "1234567890",
    "email": "info@abc.com",
    "telefon": "+90 532 123 4567",
    "adres": "İstanbul, Türkiye",
    "tip": "musteri"
}

// Yanıt (Response - 201 Created)
{
    "id": 1,
    "unvan": "ABC Ltd Şti",
    "vergi_no": "1234567890",
    "vergi_dairesi": null,
    "adres": "İstanbul, Türkiye",
    "telefon": "+90 532 123 4567",
    "email": "info@abc.com",
    "tip": "musteri",
    "aktif": true,
    "olusturma_tarihi": "2026-05-27T12:00:00+00:00",
    "guncelleme_tarihi": "2026-05-27T12:00:00+00:00"
}''')

p = doc.add_paragraph()
run = p.add_run('Örnek 2 – Kullanıcı Girişi (POST /api/v1/auth/login)')
run.bold = True
run.font.size = Pt(10)

add_code_block('''// İstek (Request Body)
{
    "email": "admin@emuhasebe.com",
    "password": "gucluSifre123"
}

// Yanıt (Response - 200 OK)
{
    "access_token": "eyJhbGciOiJIUzI1NiIs...",
    "refresh_token": "eyJhbGciOiJIUzI1NiIs...",
    "user": {
        "id": 1,
        "username": "admin",
        "email": "admin@emuhasebe.com",
        "role": "admin"
    }
}''')

p = doc.add_paragraph()
run = p.add_run('Örnek 3 – AI Analiz (POST /api/ai/analyze)')
run.bold = True
run.font.size = Pt(10)

add_code_block('''// İstek (Request Body)
{
    "prompt": "Bu ayki satış verilerini analiz et ve trend önerisi yap",
    "model": "google/gemini-2.5-flash:free"
}

// Yanıt (Response - 200 OK)
{
    "text": "Satış verileriniz incelendiğinde %15 artış trendi...",
    "model": "gemini-2.5-flash"
}''')

# ────────────────────────────
# 7.3 Önemli Kod Parçaları
# ────────────────────────────

add_heading_styled('7.3 Önemli Kod Parçaları', level=2)

doc.add_paragraph(
    'Bu bölümde projenin temel iş mantığını gösteren önemli kod parçaları sunulmaktadır. '
    'Her kod parçasının altında açıklama bulunmaktadır.'
)

# Kod Parçası 1: BaseRepository
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Kod Parçası 1 – BaseRepository (Repository Pattern)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

add_code_block('''class BaseRepository:
    """Tüm repository sınıflarının temel sınıfı."""

    def __init__(self, model, session):
        self.model = model
        self.session = session

    def _base_query(self):
        """Soft-deleted kayıtları otomatik filtreler."""
        stmt = select(self.model).where(
            getattr(self.model, 'silinme_tarihi', None) == None
        )
        return stmt

    def get_all(self, eager_load=None):
        stmt = self._base_query()
        if eager_load:
            for rel in eager_load:
                stmt = stmt.options(joinedload(getattr(self.model, rel)))
        items = list(self.session.execute(stmt).scalars().all())
        return items

    def get_by_id(self, id_):
        stmt = select(self.model).where(self.model.id == id_)
        item = self.session.execute(stmt).scalars().first()
        return item

    def create(self, **data):
        obj = self.model(**data)
        db.session.add(obj)
        db.session.flush()
        db.session.commit()
        return obj

    def delete(self, obj_or_id):
        """Soft delete: Kaydı silmez, silinme_tarihi alanını doldurur."""
        obj = obj_or_id if isinstance(obj_or_id, self.model) else self.get_by_id(obj_or_id)
        if obj is None:
            return False
        if hasattr(obj, 'silinme_tarihi'):
            obj.silinme_tarihi = datetime.now(timezone.utc)
            db.session.add(obj)
        else:
            db.session.delete(obj)
        db.session.flush()
        db.session.commit()
        return True''', "Dosya: app/repositories/base_repository.py")

p = doc.add_paragraph(
    'Açıklama: BaseRepository sınıfı, Repository Pattern tasarım kalıbını uygulayarak tüm veri erişim '
    'işlemlerini soyutlar. _base_query() metodu, soft-deleted kayıtları otomatik olarak filtreler. '
    'delete() metodu gerçek silme yerine mantıksal silme (soft delete) uygulayarak muhasebe kayıtlarının '
    'silinmeden korunmasını sağlar. Tüm repository alt sınıfları (MusteriRepository, UrunRepository vb.) '
    'bu sınıftan türetilir.'
)
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Kod Parçası 2: MusteriService
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Kod Parçası 2 – MusteriService (İş Mantığı Katmanı)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

add_code_block('''class MusteriService:
    """Müşteri iş kurallarını yöneten servis sınıfı."""

    @staticmethod
    def create(data: Dict[str, Any]) -> Musteri:
        # Vergi numarası benzersizlik kontrolü
        if data.get('vergi_no') and MusteriRepository.get_by_vergi_no(data['vergi_no']):
            raise ValueError('Vergi numarası zaten kayıtlı')
        # E-posta benzersizlik kontrolü
        if data.get('email') and MusteriRepository.get_by_email(data['email']):
            raise ValueError('Email zaten kayıtlı')

        musteri = Musteri(**data)
        db.session.add(musteri)
        db.session.commit()
        return musteri

    @staticmethod
    def update(musteri_or_id, changes: Dict[str, Any]) -> Musteri:
        if isinstance(musteri_or_id, Musteri):
            musteri = musteri_or_id
        else:
            musteri = MusteriRepository.get_by_id(musteri_or_id)
        if musteri is None:
            raise ValueError('Müşteri bulunamadı')

        # Güncelleme sırasında da benzersizlik kontrolü
        if changes.get('vergi_no'):
            existing = MusteriRepository.get_by_vergi_no(changes['vergi_no'])
            if existing and existing.id != musteri.id:
                raise ValueError('Vergi numarası başka bir müşteriye ait')

        for k, v in changes.items():
            setattr(musteri, k, v)
        db.session.add(musteri)
        db.session.commit()
        return musteri''', "Dosya: app/services/musteri_service.py")

p = doc.add_paragraph(
    'Açıklama: MusteriService sınıfı, Repository ile Controller (API Blueprint) arasında köprü görevi görür. '
    'İş kurallarını (vergi numarası benzersizliği, e-posta tekrarı kontrolü vb.) bu katmanda uygulayarak '
    'veri tutarlılığını garanti altına alır. @staticmethod dekoratörü ile stateless servis tasarımı kullanılmıştır.'
)
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Kod Parçası 3: Fatura Hesaplama Motoru
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Kod Parçası 3 – Fatura Hesaplama Motoru (Finansal Hesaplama)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

add_code_block('''class AlisFatura(db.Model, SoftDeleteMixin):
    """Alış faturaları – Numeric(12,2) hassasiyeti."""

    # Finansal alanlar: Float yerine Numeric(12,2) kullanılır
    ara_toplam    = db.Column(db.Numeric(12, 2), default=Decimal('0.00'))
    kdv_toplam    = db.Column(db.Numeric(12, 2), default=Decimal('0.00'))
    indirim_toplam = db.Column(db.Numeric(12, 2), default=Decimal('0.00'))
    genel_toplam  = db.Column(db.Numeric(12, 2), default=Decimal('0.00'))

    kalemler = db.relationship('AlisFaturaKalem', backref='fatura',
                               cascade='all, delete-orphan')

    def hesapla(self):
        """İndirimli KDV matrahı hesaplama.

        İndirim, KDV matrahını orantısal olarak düşürür.
        """
        indirim = self.indirim_toplam or Decimal('0.00')
        self.ara_toplam = sum((k.toplam for k in self.kalemler), Decimal('0.00'))

        brut_kdv = sum((k.kdv_tutar for k in self.kalemler), Decimal('0.00'))
        kdv_matrahi = max(self.ara_toplam - indirim, Decimal('0.00'))

        # KDV matrah oranı: indirim sonrası / indirim öncesi
        oran = (kdv_matrahi / self.ara_toplam) if self.ara_toplam else Decimal('0.00')
        self.kdv_toplam = brut_kdv * oran
        self.genel_toplam = kdv_matrahi + self.kdv_toplam

        if self.genel_toplam < Decimal('0.00'):
            raise ValueError(f'Geçersiz total: {self.genel_toplam}')


class AlisFaturaKalem(db.Model):
    """Fatura satır kalemi – birim fiyat × miktar × KDV hesabı."""

    miktar      = db.Column(db.Numeric(12, 2), default=Decimal('1'))
    birim_fiyat = db.Column(db.Numeric(12, 2), default=Decimal('0.00'))
    kdv_orani   = db.Column(db.Integer, default=20)
    indirim_orani = db.Column(db.Integer, default=0)

    @property
    def toplam(self):
        """Kalem toplamı (KDV hariç, indirim dahil)."""
        tutar = self.miktar * self.birim_fiyat
        indirim = tutar * (Decimal(self.indirim_orani) / Decimal('100'))
        return tutar - indirim

    @property
    def kdv_tutar(self):
        """KDV tutarı."""
        return self.toplam * (Decimal(self.kdv_orani) / Decimal('100'))''', "Dosya: app/models.py")

p = doc.add_paragraph(
    'Açıklama: Fatura hesaplama motoru, muhasebe uygulamalarında kritik öneme sahip finansal hesaplamaları yapar. '
    'Decimal tipi kullanılarak kayan nokta (floating point) hataları önlenir — bu, muhasebe yazılımlarında '
    'en önemli tasarım kararlarından biridir. hesapla() metodu, indirimli KDV matrahı hesabını orantısal olarak '
    'yapar: indirim tutarı KDV matrahını düşürür ve KDV bu düşürülmüş matrah üzerinden hesaplanır. '
    'CHECK constraint\'ler ile negatif tutar, geçersiz KDV oranı gibi durumlar veritabanı seviyesinde engellenir.'
)
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Kod Parçası 4: AI Engine
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Kod Parçası 4 – Yapay Zeka Motoru (AI Engine – Anomali Tespiti)')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

add_code_block('''/**
 * Anomali Tespiti – Z-score tabanlı anormallik tespiti.
 * Fatura tutarlarındaki istatistiksel sapmaları tespit eder.
 * @param {Array}  faturalar  - Fatura dizisi
 * @param {number} esik       - Z-score eşik değeri (varsayılan: 2.5)
 * @returns {Array} Anomali olan faturalar (z-score, seviye, sapma yönü ile)
 */
export function anomaliTespiti(faturalar, esik = 2.5) {
    if (faturalar.length < 3) return [];

    const tutarlar = faturalar
        .map(f => parseFloat(f.toplam_tutar) || parseFloat(f.genel_toplam) || 0)
        .filter(t => t > 0);

    if (tutarlar.length < 3) return [];

    // Ortalama ve standart sapma hesabı
    const ort = tutarlar.reduce((a, b) => a + b, 0) / tutarlar.length;
    const varyans = tutarlar.reduce((a, b) => a + (b - ort) ** 2, 0) / tutarlar.length;
    const stdSapma = Math.sqrt(varyans);

    if (stdSapma === 0) return [];

    return faturalar
        .filter(f => {
            const tutar = parseFloat(f.toplam_tutar) || parseFloat(f.genel_toplam) || 0;
            return tutar > 0 && Math.abs((tutar - ort) / stdSapma) > esik;
        })
        .map(f => {
            const tutar = parseFloat(f.toplam_tutar) || parseFloat(f.genel_toplam) || 0;
            const z = (tutar - ort) / stdSapma;
            return {
                ...f,
                z_skoru: Math.abs(z).toFixed(2),
                seviye: Math.abs(z) > 3.5 ? 'kritik' : 'uyari',
                sapma_yonu: tutar > ort ? 'yüksek' : 'düşük'
            };
        });
}''', "Dosya: app/static/js/ai-engine.js")

p = doc.add_paragraph(
    'Açıklama: Anomali tespiti modülü, Z-score (standart skor) istatistiksel yöntemini kullanarak '
    'fatura tutarlarındaki anormal sapmaları tespit eder. Her fatura tutarının ortalamadan kaç standart sapma '
    'uzakta olduğu hesaplanır. Eşik değer (varsayılan 2.5σ) aşılırsa fatura "anomali" olarak işaretlenir. '
    'Z-skoru 3.5σ üzerindeki sapmalar "kritik", altındakiler "uyarı" seviyesi olarak sınıflandırılır. '
    'Bu yaklaşım, sahte fatura tespiti, hatalı veri girişi ve olağandışı ticari aktivitelerin '
    'otomatik olarak tespit edilmesini sağlar.'
)
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── Kaydet ──
output_path = os.path.join(os.path.dirname(__file__), 'rapor_bolum7.docx')
doc.save(output_path)
print(f'Bolum 7 Word belgesi olusturuldu: {output_path}')
